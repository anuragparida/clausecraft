"""Build the expected redline ``.docx`` for the counterfactual demo NDA.

This script generates ``demo/expected-redline.docx`` — a Word/LibreOffice-
readable ``.docx`` showing the 5 deviations from ``known-bad-nda.pdf`` as
tracked changes (strikethroughs on the bad text, insertions of the corrected
text). The asciinema screencast downloads a redline that, when opened in
Word, should look like this file.

Why this exists
---------------
The Phase 6 spec calls for two paired artifacts:

1. ``demo/known-bad-nda.pdf`` — a hand-crafted NDA with 5 deviations
2. ``demo/expected-redline.docx`` — the expected output of running the
   system against (1) and approving all 5 deviations

The second artifact is the "spec by example" for what the system
should produce. The asciinema demos the system in action; the
expected-redline is the static reference for "this is what you
should see when you open the .docx the system generated."

How the ``.docx`` is built
--------------------------
We use ``python-docx`` for the document skeleton, then drop into
raw ``lxml`` to add ``w:ins`` / ``w:del`` tracked-changes elements
(the same pattern as ``backend/app/output/docx.py``). The
attributes match what Word / LibreOffice expect:

- ``w:author="clausecraft"``
- ``w:date="<ISO-8601 UTC timestamp>"``
- ``w:id="<sequential integer>"`` (1, 2, 3, ...)

The 5 deviations are laid out clause-by-clause. For each clause,
the original (bad) text is rendered as a deletion, and the
corrected text is rendered as an insertion. The clause heading
(\"1. Definition of Confidential Information\", etc.) is the
unmodified reference, then the strike/insert pair follows.

This is a static \"expected\" artifact, not the system output.
The HITL state machine (Phase 3 Build 3) produces a real redline
when the user clicks \"Approve\" on each of the 5 deviation
flags. The two ``.docx`` files should agree on the
strike/insert text (modulo timestamp + author), but only this
file is committed because the asciinema's reproducibility
depends on it being identical every run.
"""

from __future__ import annotations

import sys
from datetime import datetime, timezone
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# Deviation spec — 5 clauses, each with the bad text (deleted in the
# redline) and the corrected text (inserted). The corrected text is
# calibrated to the Phase 2 playbook baselines (see
# playbook/baselines/nda-en/*.yaml).
DEVIATIONS: list[dict[str, str]] = [
    {
        "id": "c1",
        "type": "definition_confidential_info",
        "heading": "1. Definition of Confidential Information.",
        "bad": (
            "\u201cConfidential Information\u201d means any and all information "
            "disclosed by one Party to the other Party, whether orally, in "
            "writing, or in any other form, without regard to whether such "
            "information is marked as confidential. The Receiving Party "
            "shall treat all such information as confidential and shall be "
            "bound by the obligations of this Agreement with respect to "
            "all such information without limitation."
        ),
        "good": (
            "\u201cConfidential Information\u201d means any non-public information "
            "disclosed by one Party (the \u201cDiscloser\u201d) to the other Party (the "
            "\u201cRecipient\u201d), whether orally, in writing, or in any other form, "
            "that is designated as confidential at the time of disclosure "
            "or that, given the nature of the information or the "
            "circumstances surrounding its disclosure, reasonably should "
            "be understood to be confidential. Confidential Information "
            "does not include information that: (a) is or becomes "
            "generally available to the public other than as a result of "
            "disclosure by the Recipient; (b) was known to the Recipient "
            "prior to its disclosure by the Discloser; (c) is or becomes "
            "available to the Recipient on a non-confidential basis from "
            "a source other than the Discloser, provided that such source "
            "is not bound by a duty of confidentiality; or (d) is "
            "required to be disclosed by law, regulation, court order, or "
            "subpoena, provided that the Recipient gives the Discloser "
            "prompt written notice of such requirement and reasonably "
            "cooperates with the Discloser in seeking a protective order."
        ),
        "rationale": (
            "Missing the 4 standard carve-outs (public, prior knowledge, "
            "independently developed, lawfully required). Without these, "
            "\u201cConfidential Information\u201d is unbounded \u2014 anything disclosed "
            "is automatically confidential."
        ),
        "category": "missing_exclusions",
    },
    {
        "id": "c2",
        "type": "term",
        "heading": "2. Term.",
        "bad": (
            "This Agreement shall commence on the Effective Date and "
            "shall continue for a period of five (5) years thereafter, "
            "unless earlier terminated in accordance with its terms. "
            "The obligations of confidentiality and non-use set forth "
            "in this Agreement shall survive the termination or "
            "expiration of this Agreement for an additional period of "
            "five (5) years, totalling ten (10) years of confidentiality "
            "obligations."
        ),
        "good": (
            "This Agreement shall commence on the Effective Date and "
            "shall continue for a period of two (2) years thereafter, "
            "unless earlier terminated in accordance with its terms. "
            "The obligations of confidentiality and non-use set forth "
            "in this Agreement shall survive the termination or "
            "expiration of this Agreement for an additional period of "
            "three (3) years; provided, however, that with respect to "
            "any Confidential Information that constitutes a trade "
            "secret under applicable law, such obligations shall survive "
            "for so long as such information remains a trade secret."
        ),
        "rationale": (
            "5 years + 5 year survival = 10 years of obligation, vs the "
            "playbook's 2 year + 3 year survival. Trade-secret protection "
            "should be \u201cfor so long as it remains a trade secret\u201d, not "
            "perpetual."
        ),
        "category": "term_too_long",
    },
    {
        "id": "c3",
        "type": "residual_knowledge",
        "heading": "3. Trade Secret Protection.",
        "bad": (
            "With respect to any information disclosed hereunder that "
            "constitutes a trade secret, the Receiving Party and its "
            "personnel shall be bound by the obligations of "
            "confidentiality and non-use set forth in this Agreement in "
            "perpetuity, without time limit, regardless of whether such "
            "information continues to qualify as a trade secret under "
            "applicable law."
        ),
        "good": (
            "Notwithstanding anything to the contrary in this Agreement, "
            "the Recipient shall be free to use any general knowledge, "
            "skills, and experience retained in the unaided memory of the "
            "Recipient\u2019s personnel in the ordinary course of their work. "
            "This Section does not grant the Recipient any license under "
            "any patent, copyright, trade secret, or other intellectual "
            "property right of the Discloser. With respect to any "
            "information disclosed hereunder that constitutes a trade "
            "secret, the obligations of confidentiality and non-use set "
            "forth in this Agreement shall survive for so long as such "
            "information remains a trade secret under applicable law."
        ),
        "rationale": (
            "No residual knowledge carve-out. The Receiving Party's "
            "employees are bound by the full confidentiality obligation "
            "even for information that remains in their unaided memory."
        ),
        "category": "missing_residual_knowledge_carveout",
    },
    {
        "id": "c4",
        "type": "governing_law",
        "heading": "4. Governing Law.",
        "bad": (
            "This Agreement shall be governed by and construed in "
            "accordance with the laws of the Cayman Islands, without "
            "giving effect to any choice-of-law or conflict-of-laws "
            "provision that would cause the laws of any other "
            "jurisdiction to apply. The Parties consent to the exclusive "
            "jurisdiction of the courts of the Cayman Islands for any "
            "action arising out of or relating to this Agreement."
        ),
        "good": (
            "This Agreement shall be governed by and construed in "
            "accordance with the laws of the State of Delaware, without "
            "giving effect to any choice-of-law or conflict-of-laws "
            "provision that would cause the laws of any other "
            "jurisdiction to apply. The Parties consent to the exclusive "
            "jurisdiction of the state and federal courts located in New "
            "Castle County, Delaware, for any action arising out of or "
            "relating to this Agreement, and irrevocably waive any "
            "objection to the venue of any such action."
        ),
        "rationale": (
            "Cayman Islands law + courts is an off-shore surprise for a "
            "US/EU counterparty. Changes the dispute-resolution venue "
            "from the playbook's Delaware / New Castle County reference."
        ),
        "category": "offshore_jurisdiction",
    },
    {
        "id": "c5",
        "type": "injunctive_relief",
        "heading": "5. Remedies.",
        "bad": (
            "The Parties agree that the sole and exclusive remedy for "
            "any breach of this Agreement shall be monetary damages, "
            "capped at a maximum aggregate liability of fifty thousand "
            "United States dollars (USD 50,000). The Parties expressly "
            "waive any right to seek injunctive relief, specific "
            "performance, or other equitable remedies in connection "
            "with this Agreement."
        ),
        "good": (
            "The Parties acknowledge and agree that any breach or "
            "threatened breach of this Agreement by the Recipient would "
            "cause irreparable harm to the Discloser for which monetary "
            "damages would be an inadequate remedy. Accordingly, the "
            "Discloser shall be entitled to seek injunctive and other "
            "equitable relief to prevent or restrain any such breach, "
            "without the necessity of proving actual damages or posting "
            "a bond, in addition to any other remedies available at law "
            "or in equity."
        ),
        "rationale": (
            "USD 50,000 cap + explicit waiver of injunctive relief "
            "removes the only remedy that actually deters a breach. "
            "Without an injunction, the Recipient can leak, pay the cap, "
            "and walk away. This is the only \u201cunacceptable\u201d (score 3) "
            "deviation \u2014 the other four are \u201cmaterial\u201d (score 2)."
        ),
        "category": "damages_cap_in_lieu_of_injunction",
    },
]


def _add_paragraph_with_tracked_changes(
    doc: Document,
    bad_text: str,
    good_text: str,
    change_id_start: int,
    author: str = "clausecraft",
) -> int:
    """Add a paragraph that contains a deletion + insertion pair.

    Returns the next change id (i.e. ``change_id_start + 2`` because
    a deletion + insertion pair uses 2 ids).
    """
    para = doc.add_paragraph()
    p_elem = para._p  # the ``w:p`` element

    # Build the w:del element (strikethrough of the bad text).
    del_elem = OxmlElement("w:del")
    del_elem.set(qn("w:id"), str(change_id_start))
    del_elem.set(qn("w:author"), author)
    del_elem.set(qn("w:date"), datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"))
    del_run = OxmlElement("w:r")
    del_run_pr = OxmlElement("w:rPr")
    del_run.append(del_run_pr)
    del_t = OxmlElement("w:delText")
    del_t.text = bad_text
    del_run.append(del_t)
    del_elem.append(del_run)
    p_elem.append(del_elem)

    # Build the w:ins element (insertion of the good text).
    ins_elem = OxmlElement("w:ins")
    ins_elem.set(qn("w:id"), str(change_id_start + 1))
    ins_elem.set(qn("w:author"), author)
    ins_elem.set(qn("w:date"), datetime.now(timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ"))
    ins_run = OxmlElement("w:r")
    ins_run_pr = OxmlElement("w:rPr")
    ins_run.append(ins_run_pr)
    ins_t = OxmlElement("w:t")
    ins_t.set(qn("xml:space"), "preserve")
    ins_t.text = good_text
    ins_run.append(ins_t)
    ins_elem.append(ins_run)
    p_elem.append(ins_elem)

    return change_id_start + 2


def build_docx(output_path: Path) -> None:
    """Build ``demo/expected-redline.docx`` from the DEVIATIONS list."""
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.core_properties.title = "Mutual NDA — Counterfactual Demo (expected redline)"
    doc.core_properties.author = "clausecraft"

    doc.add_heading("MUTUAL NON-DISCLOSURE AGREEMENT \u2014 EXPECTED REDLINE", level=1)
    doc.add_paragraph(
        "This document shows the 5 deviations the clausecraft spotter is "
        "expected to flag on demo/known-bad-nda.pdf, with the original "
        "(bad) text struck through and the corrected text inserted. "
        "Each change is annotated with the rationale + category."
    )

    next_change_id = 1
    for dev in DEVIATIONS:
        doc.add_heading(dev["heading"], level=2)
        next_change_id = _add_paragraph_with_tracked_changes(
            doc, dev["bad"], dev["good"], change_id_start=next_change_id
        )
        doc.add_paragraph(
            f"Rationale ({dev['id']}, {dev['category']}): {dev['rationale']}"
        )

    doc.save(output_path)
    print(f"Built {output_path} ({output_path.stat().st_size} bytes)")


if __name__ == "__main__":
    repo_root = Path(__file__).resolve().parent.parent
    output = repo_root / "demo" / "expected-redline.docx"
    if len(sys.argv) > 1:
        output = Path(sys.argv[1])
    build_docx(output)
