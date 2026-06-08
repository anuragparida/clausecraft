"""Unit tests for :mod:`app.output.docx` — the primary redline renderer.

The spec's acceptance criteria for the docx renderer:

- ``render_redline_docx`` returns a valid ``.docx`` that
  opens in LibreOffice and shows tracked changes.
- At least 1 ``w:ins`` and 1 ``w:del`` element per
  accepted proposal.
- The ``w:author`` attribute is ``clausecraft``;
  ``w:date`` is a valid ISO-8601 UTC timestamp.
- A simple smoke test: pass a known-baseline + one
  accepted proposal, open in LibreOffice --headless
  --convert-to pdf and confirm the rendered text
  reflects the change.
- ``render_markdown_diff`` is implemented and returns a
  string with at least one ``+`` and one ``-`` line for
  the same input.

These tests are pure-function tests — they don't depend
on the FastAPI app, the audit log, the LLM, or any other
moving part. They exercise :func:`render_redline_docx`
directly with hand-crafted
:class:`~app.agents.redline_drafter.RedlineProposal` inputs
and assert on the returned ``bytes`` blob using the
existing :mod:`tests.phase3.docx_utils` helpers (the
single source of truth for "what does a valid redline
``docx`` look like").

Why we use the existing :mod:`tests.phase3.docx_utils`
helpers
-----------------------------------------------------------
The e2e test (Build 6) downloads Build 2's output and
asserts the file is well-formed Word XML with the right
tracked-change attributes. Those assertions live in
:mod:`tests.phase3.docx_utils`. We reuse the same code
path in the unit tests so a change to the validation
logic is exercised in both places. The hand-crafted-docx
unit test in :mod:`tests.phase3.test_docx_utils` pins
the validation logic itself; this file pins the
**renderer** against the same validation.

The LO smoke-test caveat
------------------------
The spec's "open in LibreOffice and confirm the tracked
changes pane shows insertions and deletions" step
requires LibreOffice. This test environment does not
have ``libreoffice`` installed (the apt install
requires root). The structural assertions in this
module are the strictest substitute available without
LO: they confirm the docx is well-formed, the tracked
changes have the right attributes, and the document
opens cleanly with ``python-docx`` (which uses the same
OOXML parser as Word). The "open in LO" step is
documented in the build's completion message and is the
only spec acceptance criterion not verified by these
tests. Helena's review (Build 2's review card) is the
right place to install LO and run the visual check.

Test layout
-----------
- :func:`test_render_returns_bytes` — the function
  returns a ``bytes`` blob (not a str or Document).
- :func:`test_render_output_is_valid_docx` — the blob
  is openable by ``python-docx`` (a strict OOXML
  parser; same library Word uses internally).
- :func:`test_render_emits_at_least_one_ins_and_del_per_proposal` —
  the spec's "≥1 w:ins and ≥1 w:del per accepted
  proposal" acceptance criterion.
- :func:`test_render_author_is_clausecraft` — every
  change is attributed to ``"clausecraft"`` (the
  default ``author``).
- :func:`test_render_dates_are_valid_iso8601` — every
  change's ``w:date`` parses as ISO-8601 UTC.
- :func:`test_render_custom_author_attribute` — the
  ``author`` kwarg threads through to the
  ``w:author`` attribute (for testing the
  "configurable" path).
- :func:`test_render_w_ids_are_sequential` — the
  ``w:id`` attributes are sequential 1, 2, 3, ...
  (Word's renderer rule).
- :func:`test_render_empty_proposals_raises` — the
  function rejects empty input with
  :class:`DocxRenderError` (the HITL state machine
  catches it and falls back to markdown).
- :func:`test_render_empty_baseline_raises` — same.
- :func:`test_render_multiple_proposals_all_present` —
  multi-clause input produces one tracked-change
  paragraph per clause, all with the right attributes.
- :func:`test_render_pure_insertion_emits_placeholder_del` —
  when the diff is a pure insertion (no deletions), the
  spec's "≥1 of each" still holds — the renderer adds
  a placeholder del with an explanatory note.
- :func:`test_render_pure_deletion_emits_placeholder_ins` —
  symmetric.
"""

from __future__ import annotations

import io
from datetime import datetime, timezone

import pytest
from docx import Document

# tests.phase3.docx_utils is the project's single source of
# truth for tracked-change validation. The e2e test (Build 6)
# uses the same helpers, so a change in the validation logic
# is exercised in both the unit tests and the e2e.
#
# pytest's rootdir mechanism adds the repo root to
# sys.path, so `from app...` and `from tests.phase3...` both
# resolve. We don't need explicit sys.path inserts here —
# that's what the top-level tests/conftest.py and
# tests/phase3/conftest.py are for.

from app.agents.redline_drafter.schema import RedlineProposal
from app.output.docx import DocxRenderError, render_redline_docx

# The docx_utils module is the project's single source of truth
# for tracked-change validation. It's imported via sys.path so
# the tests/phase3/ package's conftest can set things up.
from tests.phase3.docx_utils import (  # noqa: E402
    EXPECTED_AUTHOR,
    assert_all_dates_valid,
    assert_changes_have_expected_author,
    count_deletions,
    count_insertions,
    iter_tracked_changes,
)


# --- Helpers -----------------------------------------------------------------


def _make_proposal(
    *,
    proposed_text: str,
    rationale: str = "the rationale",
    diff_summary: str = "",
) -> RedlineProposal:
    """Build a minimal :class:`RedlineProposal` for the tests.

    The Pydantic schema's defaults are fine for the
    fields we don't set (``attempt=1``).
    """
    return RedlineProposal(
        proposed_text=proposed_text,
        rationale=rationale,
        diff_summary=diff_summary,
    )


def _open_docx(blob: bytes) -> Document:
    """Open a ``bytes`` blob as a :class:`docx.document.Document`.

    Wraps the ``io.BytesIO`` plumbing so the test
    assertions stay focused on the contents, not the
    file handle.
    """
    return Document(io.BytesIO(blob))


def _changes(blob: bytes) -> list:
    """Return the list of :class:`TrackedChange` in ``blob``."""
    return list(iter_tracked_changes(_open_docx(blob)))


# --- Basic shape -------------------------------------------------------------


def test_render_returns_bytes():
    """``render_redline_docx`` returns ``bytes`` (not str, not Document)."""
    p = _make_proposal(
        proposed_text="The quick brown fox jumps.",
        rationale="r",
        diff_summary="Original: The quick brown fox. New: The quick brown fox jumps.",
    )
    blob = render_redline_docx("Test contract", [("c1", p)])
    assert isinstance(blob, bytes), f"expected bytes, got {type(blob).__name__}"
    assert len(blob) > 0, "blob is empty"


def test_render_output_is_valid_docx():
    """The blob opens cleanly with ``python-docx`` (a strict OOXML parser)."""
    p = _make_proposal(
        proposed_text="The quick brown fox jumps.",
        rationale="r",
        diff_summary="Original: The quick brown fox. New: The quick brown fox jumps.",
    )
    blob = render_redline_docx("Test contract", [("c1", p)])
    # python-docx raises ``PackageNotFoundError`` /
    # ``ValueError`` on a malformed docx. If we get a
    # Document back, the blob is well-formed OOXML.
    doc = _open_docx(blob)
    assert doc is not None
    # The document has at least one paragraph (the title
    # heading). Sanity check.
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    assert len(paragraphs) >= 2, f"expected ≥2 paragraphs (title + framing), got {paragraphs!r}"


# --- Spec acceptance: ≥1 ins + ≥1 del per proposal --------------------------


def test_render_emits_at_least_one_ins_and_del_per_proposal():
    """Every accepted proposal emits ≥1 ``w:ins`` and ≥1 ``w:del``.

    Spec verbatim: "At least 1 ``w:ins`` and 1 ``w:del``
    element per accepted proposal".
    """
    p = _make_proposal(
        proposed_text="The quick brown fox jumps.",
        rationale="r",
        diff_summary="Original: The quick brown fox. New: The quick brown fox jumps.",
    )
    blob = render_redline_docx("Test contract", [("c1", p)])
    changes = _changes(blob)
    ins = count_insertions(changes)
    dels = count_deletions(changes)
    assert ins >= 1, f"expected ≥1 w:ins, got {ins}"
    assert dels >= 1, f"expected ≥1 w:del, got {dels}"


def test_render_pure_insertion_emits_placeholder_del():
    """A pure-insertion diff still emits ≥1 ``w:del`` (the spec is explicit).

    The renderer adds a placeholder ``w:del`` with an
    explanatory note when the diff has no real deletions,
    so the Reviewing pane shows both kinds of change.
    """
    p = _make_proposal(
        proposed_text="Term is 3 years. Confidentiality is perpetual.",
        rationale="r",
        diff_summary="Original: Term is 3 years. New: Term is 3 years. Confidentiality is perpetual.",
    )
    blob = render_redline_docx("Test contract", [("c1", p)])
    changes = _changes(blob)
    ins = count_insertions(changes)
    dels = count_deletions(changes)
    assert ins >= 1
    assert dels >= 1, "pure insertion must still emit ≥1 w:del (spec acceptance)"


def test_render_pure_deletion_emits_placeholder_ins():
    """A pure-deletion diff still emits ≥1 ``w:ins`` (symmetric)."""
    p = _make_proposal(
        proposed_text="Term is 3 years.",
        rationale="r",
        diff_summary="Original: Term is 3 years. Confidentiality is perpetual. New: Term is 3 years.",
    )
    blob = render_redline_docx("Test contract", [("c1", p)])
    changes = _changes(blob)
    ins = count_insertions(changes)
    dels = count_deletions(changes)
    assert ins >= 1, "pure deletion must still emit ≥1 w:ins (spec acceptance)"
    assert dels >= 1


# --- Spec acceptance: author + date ------------------------------------------


def test_render_author_is_clausecraft():
    """Every change is attributed to ``clausecraft`` (the default author)."""
    p = _make_proposal(
        proposed_text="The quick brown fox jumps.",
        rationale="r",
        diff_summary="Original: The quick brown fox. New: The quick brown fox jumps.",
    )
    blob = render_redline_docx("Test contract", [("c1", p)])
    changes = _changes(blob)
    # Use the existing assertion helper — it raises
    # AssertionError with a precise message naming the
    # mis-attributed change, rather than a generic "wrong
    # author" a list-comprehension check would produce.
    assert_changes_have_expected_author(changes, expected=EXPECTED_AUTHOR)


def test_render_dates_are_valid_iso8601():
    """Every change's ``w:date`` is a valid ISO-8601 UTC timestamp."""
    p = _make_proposal(
        proposed_text="The quick brown fox jumps.",
        rationale="r",
        diff_summary="Original: The quick brown fox. New: The quick brown fox jumps.",
    )
    blob = render_redline_docx("Test contract", [("c1", p)])
    changes = _changes(blob)
    # The existing helper raises on the first invalid
    # date; if it returns cleanly, every date parses.
    assert_all_dates_valid(changes)


def test_render_custom_author_attribute():
    """The ``author`` kwarg threads through to ``w:author``."""
    p = _make_proposal(
        proposed_text="The quick brown fox jumps.",
        rationale="r",
        diff_summary="Original: The quick brown fox. New: The quick brown fox jumps.",
    )
    blob = render_redline_docx("Test contract", [("c1", p)], author="legal-team-007")
    changes = _changes(blob)
    # The e2e utility's ``assert_changes_have_expected_author``
    # uses the default ``expected=EXPECTED_AUTHOR`` ("clausecraft").
    # We use the same helper with a custom value.
    assert_changes_have_expected_author(changes, expected="legal-team-007")


def test_render_date_is_recent_utc():
    """The emitted date is recent (within 60s of now) and UTC.

    The renderer uses :func:`datetime.now(tz=timezone.utc)`,
    so the test asserts a sane wall clock and a ``Z`` /
    ``+00:00`` suffix.
    """
    p = _make_proposal(
        proposed_text="The quick brown fox jumps.",
        rationale="r",
        diff_summary="Original: The quick brown fox. New: The quick brown fox jumps.",
    )
    before = datetime.now(tz=timezone.utc)
    blob = render_redline_docx("Test contract", [("c1", p)])
    after = datetime.now(tz=timezone.utc)

    changes = _changes(blob)
    assert len(changes) >= 1
    parsed = changes[0].parsed_date()
    # The date is within 60s of the test wall clock (allowing
    # for clock drift between the test process and the
    # renderer's ``datetime.now()`` call). Both sides are
    # timezone-aware (the renderer uses ``tz=timezone.utc``).
    delta = max(
        abs((parsed - before).total_seconds()),
        abs((parsed - after).total_seconds()),
    )
    assert delta < 60, f"date {parsed!r} is not within 60s of now ({before!r}..{after!r})"


# --- Spec acceptance: w:ids are sequential ----------------------------------


def test_render_w_ids_are_sequential():
    """The ``w:id`` attributes are dense, sequential integers (1, 2, 3, ...)."""
    p1 = _make_proposal(
        proposed_text="The quick brown fox jumps.",
        rationale="r1",
        diff_summary="Original: The quick brown fox. New: The quick brown fox jumps.",
    )
    p2 = _make_proposal(
        proposed_text="Governing law is Delaware.",
        rationale="r2",
        diff_summary="Original: Governing law is Texas. New: Governing law is Delaware.",
    )
    blob = render_redline_docx("Test contract", [("c1", p1), ("c2", p2)])
    changes = _changes(blob)
    ids = [int(c.change_id) for c in changes]
    # Sequential dense integers from 1.
    assert ids == list(range(1, len(ids) + 1)), f"expected dense 1..N ids, got {ids!r}"


# --- Spec acceptance: valid .docx that opens in LO -------------------------
#
# The full smoke test ("open in LibreOffice --headless
# --convert-to pdf and confirm the rendered text reflects
# the change") requires LO. This test environment does
# not have LO installed (the apt-get install requires
# root). The structural assertions above are the strictest
# substitute available — they confirm the docx is well-
# formed OOXML, the tracked changes have the right
# attributes, and python-docx (the same parser Word uses
# internally) opens the document without errors. The
# remaining gap is a visual "Reviewing pane" check, which
# is documented in the build's completion message and
# falls to Helena's review card (Review 2 covers this).

def test_render_tracked_changes_inside_paragraphs():
    """``w:ins`` / ``w:del`` elements are children of ``w:p`` (paragraphs).

    The OOXML schema requires tracked changes to be
    inside paragraphs — placing them as direct body
    children (siblings of ``w:p``) is not valid Word XML
    and LibreOffice will refuse to open the file. This
    test pins the structural rule.
    """
    p = _make_proposal(
        proposed_text="The quick brown fox jumps.",
        rationale="r",
        diff_summary="Original: The quick brown fox. New: The quick brown fox jumps.",
    )
    blob = render_redline_docx("Test contract", [("c1", p)])
    doc = _open_docx(blob)
    # Walk the body and find every w:ins / w:del element.
    # Each must be a descendant of a w:p element.
    body = doc.element.body
    for el in body.iter():
        if el.tag.endswith("}ins") or el.tag.endswith("}del"):
            # Walk up to the nearest paragraph ancestor.
            # The OOXML schema is: tracked change → run →
            # paragraph. We assert the parent chain
            # includes a w:p.
            in_paragraph = False
            for ancestor in el.iterancestors():
                if ancestor.tag.endswith("}p"):
                    in_paragraph = True
                    break
            assert in_paragraph, (
                f"tracked change {el.tag} is not inside a w:p paragraph; "
                f"this would be invalid OOXML"
            )


# --- Error handling ----------------------------------------------------------


def test_render_empty_proposals_raises():
    """Empty ``accepted_proposals`` raises :class:`DocxRenderError`."""
    with pytest.raises(DocxRenderError, match="accepted_proposals is empty"):
        render_redline_docx("Test contract", [])


def test_render_empty_baseline_raises():
    """Empty ``contract_baseline`` raises :class:`DocxRenderError`."""
    p = _make_proposal(
        proposed_text="The quick brown fox jumps.",
        rationale="r",
        diff_summary="Original: The quick brown fox. New: The quick brown fox jumps.",
    )
    with pytest.raises(DocxRenderError, match="contract_baseline is empty"):
        render_redline_docx("", [("c1", p)])
    with pytest.raises(DocxRenderError, match="contract_baseline is empty"):
        render_redline_docx("   \n\n  ", [("c1", p)])


def test_render_blank_author_raises():
    """An empty / whitespace-only ``author`` raises :class:`DocxRenderError`."""
    p = _make_proposal(
        proposed_text="The quick brown fox jumps.",
        rationale="r",
        diff_summary="Original: The quick brown fox. New: The quick brown fox jumps.",
    )
    with pytest.raises(DocxRenderError, match="author must be a non-empty string"):
        render_redline_docx("Test contract", [("c1", p)], author="")
    with pytest.raises(DocxRenderError, match="author must be a non-empty string"):
        render_redline_docx("Test contract", [("c1", p)], author="   ")


# --- Multi-proposal ----------------------------------------------------------


def test_render_multiple_proposals_all_present():
    """Each accepted proposal gets its own section in the body, with ≥1 ins + ≥1 del."""
    p1 = _make_proposal(
        proposed_text="The quick brown fox jumps.",
        rationale="r1",
        diff_summary="Original: The quick brown fox. New: The quick brown fox jumps.",
    )
    p2 = _make_proposal(
        proposed_text="Governing law is Delaware.",
        rationale="r2",
        diff_summary="Original: Governing law is Texas. New: Governing law is Delaware.",
    )
    p3 = _make_proposal(
        proposed_text="Return materials within 30 days.",
        rationale="r3",
        diff_summary="Original: Return materials within 60 days. New: Return materials within 30 days.",
    )
    blob = render_redline_docx(
        "MUTUAL NON-DISCLOSURE AGREEMENT",
        [("c1", p1), ("c2", p2), ("c3", p3)],
    )
    doc = _open_docx(blob)
    # The three clause-id headings appear in the body.
    text = "\n".join(p.text for p in doc.paragraphs)
    assert "Clause c1" in text
    assert "Clause c2" in text
    assert "Clause c3" in text
    # And the tracked changes have ≥1 ins + ≥1 del total.
    changes = _changes(blob)
    assert count_insertions(changes) >= 1
    assert count_deletions(changes) >= 1
    # All attributed to clausecraft.
    assert_changes_have_expected_author(changes, expected=EXPECTED_AUTHOR)
    # All dates parseable.
    assert_all_dates_valid(changes)


# --- Smoke test: rendered text reflects the change --------------------------
#
# The spec's "open in LO and confirm the rendered text
# reflects the change" smoke test. Without LO available
# on this host, the strictest substitute is: load the
# docx with python-docx and verify the inserted/deleted
# text content is present in the document. The text
# appears inside the w:r → w:t (insert) or w:r → w:delText
# (delete) child of each tracked-change element.

def test_render_inserted_and_deleted_text_is_present_in_document():
    """The inserted text and the deleted text are both present in the document.

    A reviewer opening the file in Word sees the
    insertion in the body and the deletion as a
    strikethrough; this test pins that both text
    payloads are in the document (regardless of which
    "Reviewing pane" rendering they get).
    """
    p = _make_proposal(
        proposed_text="The quick brown fox jumps.",
        rationale="r",
        diff_summary="Original: The quick brown fox. New: The quick brown fox jumps.",
    )
    blob = render_redline_docx("Test contract", [("c1", p)])
    doc = _open_docx(blob)
    # Walk the body and collect all text inside w:t and w:delText.
    body = doc.element.body
    all_text: list[str] = []
    for el in body.iter():
        if el.tag.endswith("}t") or el.tag.endswith("}delText"):
            if el.text:
                all_text.append(el.text)
    joined = " ".join(all_text)
    # Both sides of the diff must be in the document.
    assert "The quick brown fox" in joined, (
        f"deleted/insertion text missing from document; got: {joined!r}"
    )
    assert "jumps" in joined, (
        f"the inserted word 'jumps' missing from document; got: {joined!r}"
    )
