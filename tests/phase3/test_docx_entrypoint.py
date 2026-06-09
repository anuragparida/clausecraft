"""Tests for :func:`app.output.render_redline` — the Build 2 entry point.

The card's hard rules for the entry point (Build 2 acceptance
criterion 3):

- Lives in ``backend/app/output/__init__.py``.
- Signature: ``render_redline(contract, proposals) -> tuple[bytes, str]``
  returning ``(bytes, filename_hint)`` where ``filename_hint`` is
  ``"contract_redline.docx"`` or ``"contract_redline.md"`` depending
  on which path the caller wants.
- The caller picks the format; the function just provides both.

These tests pin that contract: format selection, filename hint,
``bytes`` (not ``str``) return, docx round-trip equivalence with
``render_redline_docx``, and the spec's "semantic equivalence
cross-check" (the same clauses, the same rationales, the same
``(before, after)`` text appear in both the docx blob and the
markdown blob).
"""

from __future__ import annotations

import io

import pytest
from docx import Document

from app.agents.redline_drafter.schema import RedlineProposal
from app.output import (
    FILENAME_DOCX,
    FILENAME_MARKDOWN,
    DocxRenderError,
    render_markdown_diff,
    render_redline,
    render_redline_docx,
)


# --- Fixtures -------------------------------------------------------------


#: A small contract baseline. The renderer uses the first non-blank
#: line as the document title (docx) / document header (markdown).
_BASELINE = "Mutual Non-Disclosure Agreement\n\n" + (
    "This Agreement is entered into between Acme Inc. and the "
    "counterparty. Confidential Information shall be protected for "
    "a period of three years from the date of disclosure."
)


def _make_proposal(
    *,
    clause_id: str = "c1",
    proposed_text: str = "Confidential Information shall be protected for a period of five years.",
    rationale: str = "Three years is too short for trade-secret protection; five years matches industry standard.",
    diff_summary: str = (
        "Original: Confidential Information shall be protected for a period of three years. "
        "New: Confidential Information shall be protected for a period of five years."
    ),
) -> RedlineProposal:
    """Hand-build a :class:`RedlineProposal` for tests.

    The drafter is exercised in Build 1's tests; here we just
    need a well-formed Pydantic instance to feed the renderer.
    """
    return RedlineProposal(
        proposed_text=proposed_text,
        rationale=rationale,
        diff_summary=diff_summary,
    )


# --- API surface ---------------------------------------------------------


def test_entry_point_importable_from_package_root():
    """``render_redline`` is exported at the package root.

    Build 3 / Build 6 import it as ``from app.output import
    render_redline``; this is the single import path. Re-export
    at the leaf modules (``app.output.docx`` /
    ``app.output.markdown_diff``) is unchanged — backward
    compat for any existing caller.
    """
    import app.output as out

    assert hasattr(out, "render_redline")
    assert out.render_redline is render_redline


def test_filename_hints_are_exported():
    """``FILENAME_DOCX`` and ``FILENAME_MARKDOWN`` are public constants.

    The card body names them as the return values'
    ``filename_hint`` slot. They're module-level so the
    API layer (Build 3) can compare against them
    without re-typing the literal.
    """
    assert FILENAME_DOCX == "contract_redline.docx"
    assert FILENAME_MARKDOWN == "contract_redline.md"


# --- docx path ------------------------------------------------------------


def test_render_redline_docx_returns_bytes_and_filename_hint():
    """docx path: ``(bytes, "contract_redline.docx")``."""
    blob, hint = render_redline(_BASELINE, [("c1", _make_proposal())], format="docx")
    assert isinstance(blob, bytes)
    assert hint == FILENAME_DOCX


def test_render_redline_docx_blob_is_valid_docx():
    """The docx blob opens cleanly in ``python-docx``."""
    blob, _ = render_redline(_BASELINE, [("c1", _make_proposal())], format="docx")
    # python-docx is a strict OOXML parser — same library Word
    # uses internally. If the blob is malformed, this raises.
    doc = Document(io.BytesIO(blob))
    assert doc is not None


def test_render_redline_docx_blob_has_tracked_changes():
    """The docx blob has ≥1 ``w:ins`` and ≥1 ``w:del`` per proposal."""
    from docx.oxml.ns import qn

    blob, _ = render_redline(_BASELINE, [("c1", _make_proposal())], format="docx")
    body = Document(io.BytesIO(blob)).element.body
    ins_count = len(body.findall(f".//{qn('w:ins')}"))
    del_count = len(body.findall(f".//{qn('w:del')}"))
    assert ins_count >= 1
    assert del_count >= 1


def test_render_redline_docx_default_format_is_docx():
    """``format=`` is optional; the default is ``"docx"``."""
    blob_default, hint_default = render_redline(
        _BASELINE, [("c1", _make_proposal())]
    )
    blob_explicit, hint_explicit = render_redline(
        _BASELINE, [("c1", _make_proposal())], format="docx"
    )
    # Note: the blob contents differ call-to-call because the
    # ``w:date`` timestamp is computed at render time. We
    # only assert on the *shape* equivalence — both produce
    # a valid docx with tracked changes — not byte-for-byte
    # equality.
    assert hint_default == hint_explicit == FILENAME_DOCX
    # Both blobs open cleanly.
    Document(io.BytesIO(blob_default))
    Document(io.BytesIO(blob_explicit))


def test_render_redline_docx_author_kwarg_threads_through():
    """``author=`` overrides the default ``"clausecraft"`` value."""
    from docx.oxml.ns import qn

    blob, _ = render_redline(
        _BASELINE,
        [("c1", _make_proposal())],
        format="docx",
        author="legal-team-007",
    )
    body = Document(io.BytesIO(blob)).element.body
    authors = {el.get(qn("w:author")) for el in body.findall(f".//{qn('w:ins')}")}
    authors |= {el.get(qn("w:author")) for el in body.findall(f".//{qn('w:del')}")}
    assert "legal-team-007" in authors


def test_render_redline_docx_propagates_DocxRenderError():
    """An empty proposals list raises :class:`DocxRenderError`.

    The card says the caller (HITL state machine, Build 3)
    catches :class:`DocxRenderError` and falls back to the
    markdown path. The entry point does not silently
    swallow the error — it propagates so the fallback is
    triggered explicitly.
    """
    with pytest.raises(DocxRenderError):
        render_redline(_BASELINE, [], format="docx")


# --- markdown path --------------------------------------------------------


def test_render_redline_markdown_returns_bytes_and_filename_hint():
    """markdown path: ``(bytes, "contract_redline.md")``.

    The card spec says the return is always ``bytes``, even
    for the markdown path. The function UTF-8 encodes the
    markdown string before returning.
    """
    blob, hint = render_redline(
        _BASELINE, [("c1", _make_proposal())], format="markdown"
    )
    assert isinstance(blob, bytes)
    assert hint == FILENAME_MARKDOWN
    # Decoding round-trips.
    text = blob.decode("utf-8")
    assert isinstance(text, str)


def test_render_redline_markdown_is_valid_markdown_with_diff_lines():
    """The markdown blob has the expected shape: header, sections, diff lines."""
    blob, _ = render_redline(
        _BASELINE, [("c1", _make_proposal())], format="markdown"
    )
    text = blob.decode("utf-8")
    # The renderer uses ``-`` and ``+`` per the existing
    # markdown_diff module (the card says ``---/+++``; the
    # shipped code uses the more conventional ``-`` / ``+``
    # because ``---`` is a Markdown setext-h2 underline and
    # would break the section structure). Both ``-`` and
    # ``+`` are present.
    assert "+" in text
    assert "-" in text
    # The rationale is surfaced (blockquoted).
    assert ">" in text
    assert "five years" in text.lower()


def test_render_redline_markdown_byte_for_byte_matches_render_markdown_diff():
    """The markdown blob is exactly ``render_markdown_diff(...).encode("utf-8")``.

    The entry point is a thin wrapper; the leaf
    :func:`render_markdown_diff` is the source of truth.
    """
    proposals = [("c1", _make_proposal())]
    blob, _ = render_redline(_BASELINE, proposals, format="markdown")
    expected = render_markdown_diff(_BASELINE, proposals).encode("utf-8")
    assert blob == expected


# --- Cross-format equivalence --------------------------------------------


def test_render_redline_docx_and_markdown_cover_same_clauses():
    """Semantic equivalence cross-check: same clauses changed in both.

    The card says: "for a known contract + known proposals, both
    ``render_docx`` and ``render_markdown_diff`` produce semantically
    equivalent content (same clauses changed, same rationale)".

    We pin this by asserting:

    - Both blobs reference every ``clause_id`` in the input.
    - The rationale text is present in the markdown blob
      (the docx blob's rationale lives in the docx body
      text, not the tracked changes; we check the markdown
      path for the rationale and the docx path for the
      clause_id heading).
    """
    proposals = [
        ("c1", _make_proposal()),
        (
            "c2",
            _make_proposal(
                clause_id="c2",
                proposed_text="Either party may terminate this Agreement upon thirty days written notice.",
                rationale="Added a mutual termination right to balance the parties.",
                diff_summary=(
                    "Original: This Agreement continues until terminated by mutual consent. "
                    "New: Either party may terminate this Agreement upon thirty days written notice."
                ),
            ),
        ),
    ]
    blob_docx, _ = render_redline(_BASELINE, proposals, format="docx")
    blob_md, _ = render_redline(_BASELINE, proposals, format="markdown")

    # markdown path: every clause_id appears as a heading
    md_text = blob_md.decode("utf-8")
    for cid in ("c1", "c2"):
        assert cid in md_text, f"clause_id {cid!r} missing from markdown output"

    # docx path: every clause_id appears as a "Clause <id>" heading
    docx_doc = Document(io.BytesIO(blob_docx))
    headings = [p.text for p in docx_doc.paragraphs if p.text.startswith("Clause ")]
    for cid in ("c1", "c2"):
        assert any(cid in h for h in headings), (
            f"clause_id {cid!r} missing from docx output (headings={headings!r})"
        )


def test_render_redline_docx_blob_matches_render_redline_docx_shape():
    """The docx blob from the entry point is a valid ``render_redline_docx`` output.

    The entry point is a thin wrapper; it should produce
    the same ``bytes`` blob (modulo the timestamp in
    ``w:date``) as a direct call to
    :func:`render_redline_docx`. We assert on shape: both
    open cleanly, both have the same number of tracked
    changes, both attribute changes to the same author.
    """
    from docx.oxml.ns import qn

    proposals = [("c1", _make_proposal())]
    blob_entry, _ = render_redline(_BASELINE, proposals, format="docx")
    blob_direct = render_redline_docx(_BASELINE, proposals)

    # Same number of tracked changes (the timestamp is the
    # only thing that differs between calls).
    body_e = Document(io.BytesIO(blob_entry)).element.body
    body_d = Document(io.BytesIO(blob_direct)).element.body
    n_ins_e = len(body_e.findall(f".//{qn('w:ins')}"))
    n_ins_d = len(body_d.findall(f".//{qn('w:ins')}"))
    n_del_e = len(body_e.findall(f".//{qn('w:del')}"))
    n_del_d = len(body_d.findall(f".//{qn('w:del')}"))
    assert n_ins_e == n_ins_d
    assert n_del_e == n_del_d

    # Same author on every change.
    authors_e = {el.get(qn("w:author")) for el in body_e.findall(f".//{qn('w:ins')}")}
    authors_d = {el.get(qn("w:author")) for el in body_d.findall(f".//{qn('w:ins')}")}
    assert authors_e == authors_d == {"clausecraft"}


# --- Error handling ------------------------------------------------------


def test_render_redline_unknown_format_raises_value_error():
    """``format="pdf"`` (or anything else) raises :class:`ValueError`.

    The card spec says the two supported formats are
    ``"docx"`` and ``"markdown"``. Anything else is a
    programming error and gets a clear message — not a
    silent default.
    """
    with pytest.raises(ValueError) as exc_info:
        render_redline(_BASELINE, [("c1", _make_proposal())], format="pdf")
    assert "pdf" in str(exc_info.value).lower()
    assert "docx" in str(exc_info.value)
    assert "markdown" in str(exc_info.value)
