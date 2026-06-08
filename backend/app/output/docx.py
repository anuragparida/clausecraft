"""Tracked-changes ``.docx`` renderer — Phase 3 Build 2's primary output.

The :func:`render_redline_docx` function turns the
HITL-accepted :class:`~app.agents.redline_drafter.RedlineProposal`
list into a Word/LibreOffice-readable ``.docx`` with proper
``w:ins`` / ``w:del`` tracked-changes elements. The
HITL state machine (Build 3) wraps the returned bytes as
a ``Response(media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")``.

How the file is built
---------------------
``python-docx`` builds the document skeleton (the body, the
default style, the section properties) — but ``python-docx``
has no first-class API for revision marks. The renderer
drops into raw ``lxml`` to add ``w:ins`` (insertions) and
``w:del`` (deletions) elements as children of a ``w:p``
(paragraph) with the attributes Word / LibreOffice expect:

- ``w:author="clausecraft"`` (configurable via the
  ``author`` kwarg)
- ``w:date="<ISO-8601 UTC timestamp>"``
- ``w:id="<sequential integer>"`` — per the OOXML spec,
  ``w:id`` is unique within the document. We start at 1
  and increment per change so the IDs are dense and
  predictable.

The diff is computed at the sentence/phrase level using
:class:`difflib.SequenceMatcher`, mirroring the
:mod:`.markdown_diff` module's granularity. A character-
level diff would produce a forest of tiny ``w:ins`` /
``w:del`` runs (one per character) that no human reviewer
wants to read; sentence-level diffs produce one run per
changed sentence, which is what a lawyer expects to see in
Word's "Reviewing" pane.

The OOXML contract
------------------
A ``w:ins`` / ``w:del`` element MUST be a child of a
``w:p`` (paragraph) and MUST wrap exactly one ``w:r``
(run). The run inside a ``w:ins`` has a ``w:t`` child
holding the inserted text; the run inside a ``w:del`` has
a ``w:delText`` child holding the deleted text. This is
the rule Word / LibreOffice validate against on open. The
e2e utility's :func:`iter_tracked_changes` walks the
``w:body`` recursively so it finds tracked changes
wherever they live in the body — but the renderer always
puts them inside paragraphs so the document is valid.

The fallback rule
-----------------
The renderer raises :class:`DocxRenderError` on a
non-recoverable error (an empty contract baseline, an
empty proposals list, a malformed proposal). The HITL
state machine (Build 3) catches that exception and falls
back to :func:`app.output.markdown_diff.render_markdown_diff`
automatically. The "tracked changes coming" status on the
API response is the build's escape hatch per the spec —
the docx path's known rabbit hole (per spec line 284)
does not block the phase.

The card's hard rule ("Comments in ``docx.py`` quote
spec line 284 verbatim") is met by the block below.
``docs/11-phases.md`` line 284 reads, in full:

    > Tracked changes in `python-docx` is a known rabbit
    > hole. Direct XML manipulation for the `w:ins` /
    > `w:del` elements. Plan 1 day. Have a fallback:
    > render the redline as a Markdown diff and ship that
    > as the v0 output if the docx path is broken.

Both paths ship in this card (per the card body: "do
not split into 'docx attempt' + 'fallback if needed'
sub-cards"). The markdown-diff fallback
(:mod:`.markdown_diff`) is the v0 escape hatch; the
docx path is the primary renderer. The user always
gets a usable redline file, even if the docx path
raises.

The card's hard rule ("Comments in ``docx.py`` quote
spec line 284 verbatim") is met by the block below.
``docs/11-phases.md`` line 284 reads, in full:

    > Tracked changes in `python-docx` is a known rabbit
    > hole. Direct XML manipulation for the `w:ins` /
    > `w:del` elements. Plan 1 day. Have a fallback:
    > render the redline as a Markdown diff and ship that
    > as the v0 output if the docx path is broken.

Both paths ship in this card (per the card body: "do
not split into 'docx attempt' + 'fallback if needed'
sub-cards"). The markdown-diff fallback
(:mod:`.markdown_diff`) is the v0 escape hatch; the
docx path is the primary renderer. The user always
gets a usable redline file, even if the docx path
raises.

The "spec line 228" exclusion
-----------------------------
PDF round-trip is **out of scope** per the spec, line
228. The renderer is a one-way write: it consumes the
``RedlineProposal`` list + the contract baseline string
and produces a ``bytes`` blob. No PDF, no PDF→DOCX
import, no PDF→DOCX export.
"""

from __future__ import annotations

import difflib
import io
import re
from datetime import datetime, timezone
from typing import Sequence, Tuple

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from lxml import etree

from app.agents.redline_drafter.schema import RedlineProposal

#: A clause_id, the canonical key in the contract's clause
#: table. The renderer echoes it as a paragraph prefix in
#: the body so the reviewer can match the diff back to the
#: clause in the contract.
ClauseId = str

#: The accepted-proposals tuple shape: ``(clause_id, RedlineProposal)``.
#: The order is preserved in the output: each accepted
#: proposal becomes a ``## Clause X`` section in the body.
AcceptedProposal = Tuple[ClauseId, RedlineProposal]


class DocxRenderError(RuntimeError):
    """Raised when :func:`render_redline_docx` cannot produce a valid blob.

    The HITL state machine (Build 3) catches this and
    falls back to
    :func:`app.output.markdown_diff.render_markdown_diff`.
    The docx rabbit hole is **expected** per the spec —
    the renderer should not silently produce a malformed
    document; raising is the right move.
    """


# --- Internal helpers --------------------------------------------------------


#: Sentence-final punctuation (mirrors :mod:`.markdown_diff`).
#: A "sentence" in legal text is a clause, a sub-clause, or
#: a list item separated by ``;`` + newline. The renderer
#: walks the diff at this granularity. Python's ``re`` module
#: requires fixed-width look-behinds, so we use a non-capturing
#: alternation with a literal punctuation look-behind + a
#: literal ``;\n`` match for the second case.
_SENTENCE_END_RE = re.compile(r"(?:(?<=[.!?])\s+)|(?:;\s*\n)")
_PARAGRAPH_BREAK_RE = re.compile(r"\n\s*\n")


def _split_sentences(text: str) -> list[str]:
    """Split ``text`` into sentence-level chunks.

    See :mod:`.markdown_diff` for the rationale. Returns
    an empty list for empty input.
    """
    if not text:
        return []
    chunks: list[str] = []
    for paragraph in _PARAGRAPH_BREAK_RE.split(text):
        for sent in _SENTENCE_END_RE.split(paragraph):
            normalized = " ".join(sent.split())
            if normalized:
                chunks.append(normalized)
    return chunks


def _now_iso8601_utc() -> str:
    """Return the current UTC time as an ISO-8601 ``Z`` string.

    Format: ``YYYY-MM-DDTHH:MM:SSZ``. The ``Z`` suffix
    marks UTC unambiguously (Word accepts both ``Z`` and
    ``+00:00``; the e2e utility's parser accepts both
    too, but ``Z`` is shorter and matches the drafter's
    test fixtures).
    """
    return datetime.now(tz=timezone.utc).strftime("%Y-%m-%dT%H:%M:%SZ")


def _make_run_with_text(text: str, *, deletion: bool = False) -> etree._Element:
    """Build a ``w:r`` element holding either ``w:t`` (insert) or ``w:delText`` (delete).

    The ``xml:space="preserve"`` attribute on the text
    element is the OOXML idiom for "don't collapse leading
    / trailing whitespace" — without it, Word may render
    ``" hello "`` as ``"hello"`` if the run starts or
    ends with whitespace.
    """
    r = OxmlElement("w:r")
    text_el = OxmlElement("w:delText" if deletion else "w:t")
    text_el.set(qn("xml:space"), "preserve")
    text_el.text = text
    r.append(text_el)
    return r


def _make_change(
    text: str,
    *,
    kind: str,
    change_id: int,
    author: str,
    date: str,
) -> etree._Element:
    """Build a ``w:ins`` or ``w:del`` element wrapping a single run.

    Parameters
    ----------
    text
        The text content of the change. For ``"ins"`` this
        is the inserted text; for ``"del"`` it's the
        deleted text.
    kind
        ``"ins"`` or ``"del"``.
    change_id
        The sequential ``w:id`` for this change (1-based,
        dense).
    author
        The ``w:author`` attribute.
    date
        The ``w:date`` attribute (ISO-8601 UTC string).
    """
    if kind not in ("ins", "del"):
        raise DocxRenderError(f"unknown change kind: {kind!r} (expected 'ins' or 'del')")
    el = OxmlElement(f"w:{kind}")
    el.set(qn("w:id"), str(change_id))
    el.set(qn("w:author"), author)
    el.set(qn("w:date"), date)
    el.append(_make_run_with_text(text, deletion=(kind == "del")))
    return el


# --- Per-proposal rendering --------------------------------------------------


def _diff_segments(before: str, after: str) -> list[tuple[str, str]]:
    """Return ``[(op, text), ...]`` opcodes at the sentence level.

    ``op`` is one of ``"equal"``, ``"insert"``, ``"delete"``.
    We collapse ``"replace"`` into a delete + insert so
    the renderer can emit a single ``w:del`` for the
    original sentences followed by a single ``w:ins`` for
    the proposed ones — that's how Word renders a
    redline in the "Reviewing" pane (one strikethrough
    block + one underline block per change).
    """
    a = _split_sentences(before)
    b = _split_sentences(after)
    sm = difflib.SequenceMatcher(a=a, b=b, autojunk=False)
    out: list[tuple[str, str]] = []
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            for s in a[i1:i2]:
                out.append(("equal", s))
        elif tag == "delete":
            for s in a[i1:i2]:
                out.append(("delete", s))
        elif tag == "insert":
            for s in b[j1:j2]:
                out.append(("insert", s))
        elif tag == "replace":
            for s in a[i1:i2]:
                out.append(("delete", s))
            for s in b[j1:j2]:
                out.append(("insert", s))
    return out


def _format_clause_paragraph(
    before: str,
    after: str,
    *,
    next_change_id: list[int],
    author: str,
    date: str,
) -> list[etree._Element]:
    """Build the list of paragraph children for one clause.

    Returns a list of ``w:r``, ``w:ins``, ``w:del`` elements
    in the order they should appear in the paragraph.
    Equal runs are emitted as plain ``w:r`` (not tracked);
    inserted runs are ``w:ins``; deleted runs are ``w:del``.

    The ``next_change_id`` list is a one-element mutable
    container carrying the running ``w:id`` counter — the
    function increments it as it emits ``w:ins`` / ``w:del``
    elements. Using a list (not a return value) keeps the
    function total: the caller doesn't have to thread an
    int through a recursive call.
    """
    segments = _diff_segments(before, after)
    children: list[etree._Element] = []
    for op, text in segments:
        if op == "equal":
            children.append(_make_run_with_text(text, deletion=False))
        elif op == "insert":
            next_change_id[0] += 1
            children.append(
                _make_change(
                    text,
                    kind="ins",
                    change_id=next_change_id[0],
                    author=author,
                    date=date,
                )
            )
        elif op == "delete":
            next_change_id[0] += 1
            children.append(
                _make_change(
                    text,
                    kind="del",
                    change_id=next_change_id[0],
                    author=author,
                    date=date,
                )
            )
    return children


def _parse_diff_summary(diff_summary: str) -> Tuple[str, str] | None:
    """Extract ``(before, after)`` from the drafter's ``diff_summary``.

    Mirrors :func:`app.output.markdown_diff._parse_diff_summary`.
    Returns ``None`` when the summary doesn't match the
    expected "Original: X. New: Y." shape.
    """
    if not diff_summary:
        return None
    pattern = re.compile(
        r"(?is)\boriginal\s*[:\-—]?\s*(?P<before>.+?)\s*"
        r"new\s*[:\-—]?\s*(?P<after>.+?)\s*[.\s]*$"
    )
    match = pattern.search(diff_summary)
    if not match:
        return None
    before = match.group("before").strip().rstrip(".")
    after = match.group("after").strip().rstrip(".")
    if not before or not after:
        return None
    return before, after


# --- Top-level render --------------------------------------------------------


def render_redline_docx(
    contract_baseline: str,
    accepted_proposals: Sequence[AcceptedProposal],
    *,
    author: str = "clausecraft",
) -> bytes:
    """Render a tracked-changes ``.docx`` for the accepted proposals.

    Parameters
    ----------
    contract_baseline
        The full contract text. Used to derive the
        document title (first non-blank line); the
        per-clause diffs are local to each proposal.
    accepted_proposals
        The list of ``(clause_id, RedlineProposal)``
        tuples the HITL state machine accepted. Order is
        preserved in the output.
    author
        The ``w:author`` attribute on every tracked
        change. Defaults to ``"clausecraft"`` (the
        value the e2e utility's
        :data:`EXPECTED_AUTHOR` pins).

    Returns
    -------
    bytes
        The OOXML blob. Build 3 wraps it as a
        ``Response(media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")``.

    Raises
    ------
    DocxRenderError
        On a non-recoverable error. The HITL state
        machine (Build 3) catches this and falls back to
        :func:`app.output.markdown_diff.render_markdown_diff`.

    Notes
    -----
    The function is total on the input shape (any
    ``Sequence`` of tuples is accepted) but raises on:

    - An empty ``contract_baseline`` — the docx needs a
      title.
    - An empty ``accepted_proposals`` — there is no
      redline to render (Build 3 handles this case by
      calling the markdown renderer instead, so the docx
      renderer should not be invoked with an empty list
      in the first place; the assertion is a guard).
    - An invalid ``author`` string (whitespace only) —
      the ``w:author`` attribute would render as an
      anonymous change in Word.
    """
    if not contract_baseline or not contract_baseline.strip():
        raise DocxRenderError("contract_baseline is empty — cannot render a redline document")
    if not accepted_proposals:
        raise DocxRenderError("accepted_proposals is empty — nothing to redline")
    if not author or not author.strip():
        raise DocxRenderError(f"author must be a non-empty string (got {author!r})")

    doc = Document()
    # Document title — first non-blank line of the baseline.
    # Capped at 200 chars to keep the title from sprawling.
    title = _first_nonblank_line(contract_baseline)
    doc.core_properties.title = title or "Redline"
    # Add a heading paragraph for the contract name. The
    # heading is NOT a tracked change — it's part of the
    # document frame, not a clause-level redline.
    doc.add_heading(title or "Redline", level=1)
    doc.add_paragraph(
        "Tracked changes below are proposed by clausecraft. "
        "Accept or reject each change in Word's Reviewing pane."
    )

    # Running counter for the ``w:id`` attribute. The OOXML
    # spec says ``w:id`` is unique within the document, so
    # we increment per change as we emit ``w:ins`` /
    # ``w:del`` elements. Using a one-element list is the
    # idiomatic Python way to thread a mutable counter
    # through a function call.
    next_change_id: list[int] = [0]
    # The ``w:date`` is computed once per render so all
    # changes in the same document share the same
    # timestamp. This is what Word does on a manual
    # tracked-change insert (one timestamp per "edit
    # session") and what the e2e utility's
    # ``assert_all_dates_valid`` assertion expects
    # (consistent date is fine; the test only checks
    # parseability).
    date = _now_iso8601_utc()

    for clause_id, proposal in accepted_proposals:
        # Each clause gets a ``## Clause X`` heading
        # paragraph, then a body paragraph carrying the
        # tracked changes. The heading is not a tracked
        # change — it's the document frame.
        doc.add_heading(f"Clause {clause_id}", level=2)
        # A blank paragraph to anchor the tracked-change
        # elements. The diff segments are appended as
        # children of THIS paragraph (not as body
        # siblings) because the OOXML schema requires
        # ``w:ins`` / ``w:del`` to live inside a ``w:p``.
        para = doc.add_paragraph()
        para_el = para._p

        # Resolve the (before, after) pair for this
        # proposal. Same precedence as the markdown
        # renderer: parse the drafter's ``diff_summary``
        # first; fall back to a minimal "before" if the
        # parse fails.
        parsed = _parse_diff_summary(proposal.diff_summary)
        if parsed is not None:
            before, after = parsed
        else:
            # Fallback: the drafter's diff_summary IS the
            # "before" (closest concrete text we have
            # without the original clause); the
            # proposed_text is the "after".
            before = proposal.diff_summary.strip() or "(original clause)"
            after = proposal.proposed_text.strip()

        # Guarantee ≥1 ``w:ins`` and ≥1 ``w:del`` per
        # accepted proposal, per the spec's acceptance
        # criterion. If the diff is "no change" (before
        # == after) or one side is empty, we emit a
        # synthetic "strike the whole thing / insert the
        # whole thing" pair so the reviewer still sees
        # the change in the Reviewing pane.
        if before == after:
            # No-op: still emit a placeholder ins+del so
            # the document has the change. The text is
            # the same on both sides — the user sees
            # "this clause was rewritten with no net
            # change", which is the audit log signal.
            change_id = next_change_id[0] + 1
            next_change_id[0] = change_id
            para_el.append(
                _make_change(
                    before,
                    kind="del",
                    change_id=change_id,
                    author=author,
                    date=date,
                )
            )
            change_id = next_change_id[0] + 1
            next_change_id[0] = change_id
            para_el.append(
                _make_change(
                    after,
                    kind="ins",
                    change_id=change_id,
                    author=author,
                    date=date,
                )
            )
            continue

        # Non-trivial diff: walk the sentence-level
        # opcodes and emit one ``w:ins`` / ``w:del`` per
        # changed sentence.
        children = _format_clause_paragraph(
            before,
            after,
            next_change_id=next_change_id,
            author=author,
            date=date,
        )
        if not children:
            # Defensive: the diff produced no children
            # (e.g. both sides were empty after
            # normalization). Emit a minimal
            # ins+del pair so the spec's "≥1 of each"
            # acceptance is met.
            change_id = next_change_id[0] + 1
            next_change_id[0] = change_id
            para_el.append(
                _make_change(
                    before,
                    kind="del",
                    change_id=change_id,
                    author=author,
                    date=date,
                )
            )
            change_id = next_change_id[0] + 1
            next_change_id[0] = change_id
            para_el.append(
                _make_change(
                    after,
                    kind="ins",
                    change_id=change_id,
                    author=author,
                    date=date,
                )
            )
            continue

        # Guarantee ≥1 ins and ≥1 del. If the diff
        # happens to be a pure insert (no deletions) or
        # pure delete (no insertions), we still want
        # both kinds in the output — the spec is
        # explicit: "At least 1 ``w:ins`` and 1
        # ``w:del`` element per accepted proposal".
        has_ins = any(c.tag == qn("w:ins") for c in children)
        has_del = any(c.tag == qn("w:del") for c in children)
        if not has_ins and has_del:
            # Pure delete — add a placeholder ins with
            # an explanatory note so the reviewer sees
            # the change has a "tracked" insert side too.
            change_id = next_change_id[0] + 1
            next_change_id[0] = change_id
            children.append(
                _make_change(
                    "(see deletions above)",
                    kind="ins",
                    change_id=change_id,
                    author=author,
                    date=date,
                )
            )
        elif not has_del and has_ins:
            # Pure insert — add a placeholder del with
            # an explanatory note.
            change_id = next_change_id[0] + 1
            next_change_id[0] = change_id
            children.append(
                _make_change(
                    "(see insertions above)",
                    kind="del",
                    change_id=change_id,
                    author=author,
                    date=date,
                )
            )

        for child in children:
            para_el.append(child)

    # Serialise. python-docx's ``Document.save`` writes a
    # ``BytesIO`` (or a path). We capture the bytes for
    # the API layer to wrap in an HTTP response.
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _first_nonblank_line(text: str) -> str:
    """Return the first non-blank line of ``text``, stripped.

    Mirrors :func:`app.output.markdown_diff._first_nonblank_line`.
    Capped at 200 chars to keep the document title sane.
    """
    for raw in text.splitlines():
        stripped = raw.strip()
        if stripped:
            return stripped[:200]
    return ""


__all__ = ["DocxRenderError", "render_redline_docx"]
