"""DOCX ingest — Phase 1.

Uses python-docx to extract paragraphs from a .docx in document order.
Headings (``Heading 1``/``Heading 2``) are surfaced as ``section``
metadata on each paragraph; everything else is treated as body text.

The function returns a :class:`DocxDocument`. The shape mirrors
:class:`app.ingest.pdf.PdfDocument` closely enough that the downstream
chunker can treat the two as the same kind of "text + section metadata"
object — but the dataclasses are kept separate because DOCX does not
have pages in the PDF sense.
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field

from docx import Document  # type: ignore[import-not-found]
from docx.oxml.ns import qn  # type: ignore[import-not-found]

logger = logging.getLogger(__name__)


@dataclass
class DocxParagraph:
    """A single paragraph with optional heading context."""

    paragraph_index: int  # 0-based, document order
    text: str
    style: str  # e.g. "Normal", "Heading 1", "Heading 2", "Title"
    is_heading: bool
    heading_level: int  # 0 for non-headings, else 1..9


@dataclass
class DocxDocument:
    """Result of opening + extracting from a .docx."""

    paragraphs: list[DocxParagraph] = field(default_factory=list)
    full_text: str = ""
    char_count: int = 0
    detected_scan: bool = False  # DOCX is never "scanned" — always False
    scanned_warning: str = ""


def _heading_level(style_name: str) -> int:
    """Map a python-docx style name to an integer heading level.

    ``Heading 1`` → 1, ``Heading 2`` → 2, etc. Anything else (including
    ``Title`` and ``Normal``) → 0. Case-insensitive; tolerates trailing
    whitespace and the "HeadingN" no-space variant some Word versions emit.
    """
    if not style_name:
        return 0
    name = style_name.strip()
    if not name:
        return 0
    # python-docx writes styles like "Heading 1"; some DOCX producers
    # collapse the space. Handle both.
    if name.lower().startswith("heading "):
        try:
            return int(name.split(" ", 1)[1])
        except (ValueError, IndexError):
            return 0
    if name.lower().startswith("heading"):
        digits = "".join(ch for ch in name if ch.isdigit())
        return int(digits) if digits else 0
    return 0


def extract_docx(data: bytes) -> DocxDocument:
    """Open ``data`` (a .docx byte string) and return its paragraphs."""
    try:
        doc = Document(data)  # python-docx accepts a file-like
    except Exception as exc:  # noqa: BLE001
        raise ValueError(f"DOCX could not be opened: {exc}") from exc

    paragraphs: list[DocxParagraph] = []
    for idx, para in enumerate(doc.paragraphs):
        text = para.text or ""
        style = para.style.name if para.style is not None else "Normal"
        level = _heading_level(style)
        paragraphs.append(
            DocxParagraph(
                paragraph_index=idx,
                text=text,
                style=style,
                is_heading=level > 0 or style.strip().lower() == "title",
                heading_level=level,
            )
        )

    full_text = "\n\n".join(p.text for p in paragraphs)
    return DocxDocument(
        paragraphs=paragraphs,
        full_text=full_text,
        char_count=len(full_text),
        detected_scan=False,
        scanned_warning="",
    )


# Re-export the qn helper for downstream modules that want to inspect
# raw OOXML — not used in Phase 1, but keeps the import available.
_ = qn
